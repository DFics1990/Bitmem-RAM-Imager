import os
import subprocess
import threading
import customtkinter as ctk
from tkinter import filedialog, messagebox, StringVar
import hashlib
from datetime import datetime
import time
import psutil
from fpdf import FPDF
from PIL import Image
from customtkinter import CTkImage
from pathlib import Path
import json
from docx import Document
import html
import platform
import logging
import zipfile
import sys
import zipfile
import shutil
import webbrowser
import winreg
from threading import Thread


os.environ['PYTHONUTF8'] = '1'
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('bitmem.log'),
        logging.StreamHandler()
    ]
)

# Configure appearance
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("dark-blue")

# Color Theme
DARK_BACKGROUND = "#1E2A38"
DARK_PRIMARY = "#3498DB"
DARK_SECONDARY = "#2C3E50"
DARK_ACCENT = "#2980B9"
DARK_TEXT = "#ECF0F1"

LIGHT_BACKGROUND = "#FFFFFF"
LIGHT_PRIMARY = "#3498DB"
LIGHT_SECONDARY = "#F4F4F4"
LIGHT_ACCENT = "#2980B9"
LIGHT_TEXT = "#2C3E50"

# Global color variables
BACKGROUND_COLOR = DARK_BACKGROUND
PRIMARY_COLOR = DARK_PRIMARY
SECONDARY_COLOR = DARK_SECONDARY
ACCENT_COLOR = DARK_ACCENT
TEXT_COLOR = DARK_TEXT

# Base directory
BASE_DIR = Path(__file__).parent
ASSETS_DIR = BASE_DIR / "assets"
ICONS_DIR = ASSETS_DIR / "icons"
IMAGES_DIR = ASSETS_DIR / "images"
CONFIG_FILE = BASE_DIR / "bitmem_config.json"
VOLATILITY_PATH = BASE_DIR / "volatility3"
MAX_RECENT_CASES = 5

# Add these with your other constants
SYMBOLS_DIR = BASE_DIR / "symbols"
CACHE_DIR = BASE_DIR / "volatility_cache"

# Volatility plugins with time estimates (in seconds for 8GB RAM)
ALL_VOLATILITY_PLUGINS = {
    "windows.pslist.PsList": {"time": 5, "slow": False},
    "windows.psscan.PsScan": {"time": 10, "slow": False},
    "windows.malfind.Malfind": {"time": 30, "slow": True},
    "windows.registry.printkey.PrintKey": {"time": 15, "slow": False},
    "windows.netscan.NetScan": {"time": 90, "slow": True},
    "windows.info.Info": {"time": 5, "slow": False},
  }

def load_config():
    """Load configuration from file"""
    default_config = {
        'recent_cases': [],
        'volatility_plugins': ["windows.pslist.PsList", "windows.pstree.PsTree"],
        'last_session': {},
        'appearance_mode': "Dark"
    }
    
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, 'r') as f:
                content = f.read()
                if not content.strip():  # If file is empty
                    return default_config
                config = json.loads(content)
                return {**default_config, **config}  # Merge with defaults
        except Exception as e:
            logging.error(f"Error loading config: {e}")
            # Create fresh config file
            with open(CONFIG_FILE, 'w') as f:
                json.dump(default_config, f, indent=4)
            return default_config
    return default_config

def save_config(config):
    """Save configuration to file"""
    try:
        with open(CONFIG_FILE, 'w') as f:
            json.dump(config, f, indent=4)
    except Exception as e:
        logging.error(f"Error saving config: {e}")

def show_error(msg):
    """Show error message"""
    logging.error(msg)
    messagebox.showerror("Error", msg)

def show_info(msg):
    """Show info message"""
    logging.info(msg)
    messagebox.showinfo("Information", msg)

def check_python_installed():
    """Check if Python is installed and in PATH."""
    try:
        result = subprocess.run(
            ["python", "--version"],
            capture_output=True,
            text=True,
            creationflags=subprocess.CREATE_NO_WINDOW
        )
        return "Python" in result.stdout or "Python" in result.stderr
    except:
        return False

def calculate_hash(file_path, algorithm="sha256"):
    """Calculate file hash"""
    hash_func = hashlib.new(algorithm)
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_func.update(chunk)
    return hash_func.hexdigest()

def verify_file_integrity(original_path, copied_path, algorithm="sha256"):
    """Verify two files have identical hashes"""
    original_hash = calculate_hash(original_path, algorithm)
    copied_hash = calculate_hash(copied_path, algorithm)
    return original_hash == copied_hash



def generate_filename(case_name, case_id, memory_id, extension=".raw"):
    """Generate standardized filename"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_case_name = "".join(c for c in case_name if c.isalnum() or c in (' ', '_')).rstrip()
    return f"Case_{safe_case_name}_ID{case_id}_Mem{memory_id}_{timestamp}{extension}"

def generate_report(output_file, investigator_name, investigator_id, case_name, case_id, memory_id, hash_value, image_size, elapsed_time, report_format, hash_algorithm, **kwargs):
    """Generate report in selected format"""
    report_file_name = f"Case-{case_name.replace(' ', '_')}_Report.{report_format}"
    report_file = os.path.join(os.path.dirname(output_file), report_file_name)
    
    capture_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    minutes = int(elapsed_time // 60)
    seconds = elapsed_time % 60
    duration_str = f"{minutes} minutes {seconds:.2f} seconds"
    formatted_image_size = "{:,}".format(image_size)
    
    report_data = [
        ("Investigator Name:", investigator_name),
        ("Investigator ID:", investigator_id),
        ("Case Name:", case_name),
        ("Case ID:", case_id),
        ("Memory ID:", memory_id),
        ("Capture Time:", capture_time),
        ("Image Size:", f"{formatted_image_size} bytes ({image_size/1024/1024:,.2f} MB)"),
        ("Hash Algorithm:", hash_algorithm),
        ("Hash Value:", hash_value),
        ("Image Location:", output_file.replace('\\', '/')),  # Normalize path
        ("Capture Duration:", duration_str),
    ]
    if "volatility_time" in kwargs:
       report_data.append(("Volatility Analysis Time:", kwargs["volatility_time"]))
    if "compression_time" in kwargs:
       report_data.append(("Compression Time:", kwargs["compression_time"]))

    if report_format.lower() == "txt":
        with open(report_file, "w", encoding='utf-8') as f:
            f.write("=== MEMORY CAPTURE REPORT ===\n\n")
            for key, value in report_data:
                f.write(f"{key} {value}\n")
    
    elif report_format.lower() == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "MEMORY CAPTURE REPORT", ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", size=12)
        
        for key, value in report_data:
            # Handle long paths by splitting them
            if key == "Image Location:":
                path_parts = [value[i:i+70] for i in range(0, len(value), 70)]
                pdf.cell(0, 10, f"{key} {path_parts[0]}", ln=True)
                for part in path_parts[1:]:
                    pdf.cell(0, 10, f"            {part}", ln=True)
            else:
                pdf.cell(0, 10, f"{key} {value}", ln=True)
            pdf.ln(5)
        
        pdf.output(report_file)
    
    elif report_format.lower() == "html":
        with open(report_file, "w", encoding='utf-8') as f:
            f.write("<!DOCTYPE html>\n<html>\n<head>\n<title>Memory Capture Report</title>\n")
            f.write("<style>\nbody {font-family: Arial; margin: 20px; line-height: 1.6;}\n")
            f.write("h1 {color: #3498DB; text-align: center;}\n")
            f.write("table {border-collapse: collapse; width: 100%; margin-bottom: 20px;}\n")
            f.write("th, td {border: 1px solid #ddd; padding: 8px; text-align: left;}\n")
            f.write("th {background-color: #f2f2f2;}\n</style>\n</head>\n<body>\n")
            f.write("<h1>Memory Capture Report</h1>\n<table>\n")
            
            for key, value in report_data:
                # Handle long paths by splitting them
                if key == "Image Location:":
                    path_parts = [value[i:i+80] for i in range(0, len(value), 80)]
                    f.write(f"<tr><th>{key}</th><td>{html.escape(str(path_parts[0]))}</td></tr>\n")
                    for part in path_parts[1:]:
                        f.write(f"<tr><th></th><td>{html.escape(str(part))}</td></tr>\n")
                else:
                    f.write(f"<tr><th>{key}</th><td>{html.escape(str(value))}</td></tr>\n")
            
            f.write("</table>\n</body>\n</html>")
    
    elif report_format.lower() == "docx":
        doc = Document()
        doc.add_heading('Memory Capture Report', 0)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Value'
        
        for key, value in report_data:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            # Handle long paths by splitting them
            if key == "Image Location:":
                path_parts = [value[i:i+80] for i in range(0, len(value), 80)]
                row_cells[1].text = path_parts[0]
                for part in path_parts[1:]:
                    row_cells = table.add_row().cells
                    row_cells[1].text = part
            else:
                row_cells[1].text = str(value)
        
        doc.save(report_file)
    
    elif report_format.lower() == "json":
        report_dict = {k.strip(':'): v for k, v in report_data}
        with open(report_file, "w", encoding='utf-8') as f:
            json.dump(report_dict, f, indent=4, ensure_ascii=False)
    
    logging.info(f"Report generated at {report_file}")
    return report_file


def install_imdisk():
    try:
        # Download or use bundled installer
        installer_path = os.path.join("tools", "imdiskinst.exe")
        
        subprocess.run(
            [installer_path, "/S"],  # Silent install
            check=True,
            creationflags=subprocess.CREATE_NO_WINDOW
        )
        return True
    except Exception as e:
        print(f"ImDisk install failed: {e}")
        return False


class MemoryCaptureApp:
    def __init__(self, root):
        self.root = root
        self.root.title("BitMem")
        self.root.geometry("900x700")
        
        self.user_data = {}
        self.output_file = ""
        self.capture_active = False
        self.capture_thread = None
        self.capture_process = None
        self.capture_start_time = None
        self.capture_bytes_copied = 0
        self.total_memory_size = psutil.virtual_memory().total
        self.hash_algorithm = StringVar(value="sha256")
        self.resume_capture = False
        self.partial_capture_path = None
        self.current_step = 1
        
        self.analysis_time = None
        self.compression_time = None

        self.enable_compression = False
        self.enable_volatility = False
        self.enable_preview = False
        
        self.all_plugins = ALL_VOLATILITY_PLUGINS
        self.config = load_config()
        self.selected_plugins = self.config.get('volatility_plugins', ["windows.pslist.PsList", "windows.pstree.PsTree"])
        self.live_analysis_active = False
        
        # New variables for finalizing progress tracking
        self.finalizing_steps = []
        self.current_finalizing_step = 0
        self.finalizing_start_time = None
        self.finalizing_progress_bar = None
        self.finalizing_progress_label = None
        self.finalizing_time_label = None
        self.finalizing_eta_label = None
        
        self.load_icons()
        
        icon_path = ASSETS_DIR / "newbitmemico.ico"
        if icon_path.exists():
            try:
                self.root.iconbitmap(str(icon_path))
            except Exception as e:
                logging.error(f"Error setting window icon: {e}")

        background_image_path = IMAGES_DIR / "image.jpg"
        if background_image_path.exists():
            try:
                self.bg_image = Image.open(background_image_path)
                self.background_photo = CTkImage(light_image=self.bg_image, 
                                               dark_image=self.bg_image,
                                               size=(self.root.winfo_screenwidth(), 
                                                     self.root.winfo_screenheight()))
                self.background_label = ctk.CTkLabel(root, image=self.background_photo, text="")
                self.background_label.place(x=0, y=0, relwidth=1, relheight=1)
            except Exception as e:
                logging.error(f"Error loading background image: {e}")

        self.main_container = ctk.CTkFrame(root, fg_color=SECONDARY_COLOR, corner_radius=15)
        self.main_container.place(relx=0.5, rely=0.5, anchor="center", relwidth=0.5, relheight=0.99)
        
        self.show_welcome_screen()
    




    def _format_path(self, path):
        """Format path to display with line breaks if too long"""
        max_length = 60
        if len(path) <= max_length:
            return path
        
        parts = []
        while path:
            parts.append(path[:max_length])
            path = path[max_length:]
        return "\n".join(parts)
    
    def load_icons(self):
        self.icons = {}
        icon_files = {
            'investigator_name': "invname.png",
            'investigator_id': "invid.png",
            'case_name': "casename.png",
            'case_id': "caseid.png",
            'memory_id': "memoryid.png",
            'report_format': "report format.png",
            'hash': "hash.png",
            'recent': "recent.png",
            'history': "history.png"
        }
        
        for key, filename in icon_files.items():
            icon_path = ICONS_DIR / filename
            if icon_path.exists():
                try:
                    img = Image.open(icon_path)
                    img = img.resize((20, 20), Image.Resampling.LANCZOS)
                    self.icons[key] = CTkImage(light_image=img, dark_image=img, size=(20, 20))
                except Exception as e:
                    logging.error(f"Error loading {key} icon: {e}")
    
    def clear_container(self):
        for widget in self.main_container.winfo_children():
            widget.destroy()
    
    def show_welcome_screen(self):
        self.current_step = 0
        self.clear_container()
        
        container = ctk.CTkFrame(self.main_container, fg_color="transparent")
        container.pack(expand=True, fill="both", padx=20, pady=20)
        
        title_frame = ctk.CTkFrame(container, fg_color="transparent")
        title_frame.pack(pady=(60, 20))
        
        self.welcome_label = ctk.CTkLabel(
            title_frame,
            text="Welcome!",
            font=("Segoe UI", 14),
            text_color=TEXT_COLOR
        )
        self.welcome_label.pack()
        
        self.line1 = ctk.CTkLabel(
            title_frame,
            text="Capture RAM Image with",
            font=("Segoe UI", 24, "bold"),
            text_color=PRIMARY_COLOR
        )
        self.line1.pack()
        
        self.line2 = ctk.CTkLabel(
            title_frame,
            text="BitMem",
            font=("Segoe UI", 36, "bold"),
            text_color=PRIMARY_COLOR
        )
        self.line2.pack(pady=5)
        
        self.start_btn = ctk.CTkButton(
            container,
            text="Start",
            font=("Segoe UI", 18, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            width=120,
            height=120,
            corner_radius=60,
            command=self.show_step1
        )
        self.start_btn.pack(pady=40)
        
        # Add Last Session button if we have last session data
        if self.config.get('last_session', {}):
            last_session_btn = ctk.CTkButton(
                container,
                text="View Last Session",
                font=("Segoe UI", 15),
                fg_color=SECONDARY_COLOR,
                hover_color=ACCENT_COLOR,
                command=self.show_last_session
            )
            last_session_btn.pack(pady=10)
        
        self.thanks_label = ctk.CTkLabel(
            container,
            text="Thanks for using BitMem!",
            font=("Segoe UI", 14),
            text_color=TEXT_COLOR
        )
        self.thanks_label.pack()
        
        self.add_utility_buttons()
    
    def show_last_session(self):
        """Show details of the last session"""
        last_session = self.config.get('last_session', {})
        if not last_session:
            show_info("No previous session data found")
            return
        
        dialog = ctk.CTkToplevel(self.root)
        dialog.title("Last Session Details")
        dialog.geometry("600x400")
        dialog.transient(self.root)
        dialog.grab_set()
        
        container = ctk.CTkFrame(dialog)
        container.pack(expand=True, fill="both", padx=10, pady=10)
        
        # Title with icon
        title_frame = ctk.CTkFrame(container, fg_color="transparent")
        title_frame.pack(pady=10)
        
        if 'history' in self.icons:
            icon_label = ctk.CTkLabel(title_frame, image=self.icons['history'], text="")
            icon_label.pack(side="left", padx=(0, 10))
        
        ctk.CTkLabel(
            title_frame,
            text="Last Session Details",
            font=("Segoe UI", 16, "bold"),
            text_color=PRIMARY_COLOR
        ).pack(side="left")
        
        # Scrollable content
        scroll_frame = ctk.CTkScrollableFrame(container)
        scroll_frame.pack(expand=True, fill="both", padx=10, pady=10)
        
        # Display session details
        details = [
            ("Case Name:", last_session.get('case_name', 'N/A')),
            ("Investigator:", f"{last_session.get('investigator_name', 'N/A')} (ID: {last_session.get('investigator_id', 'N/A')})"),
            ("Memory ID:", last_session.get('memory_id', 'N/A')),
            ("Date:", last_session.get('timestamp', 'N/A')),
            ("Options:", ", ".join([
                f"Compression: {'Yes' if last_session.get('compression', False) else 'No'}",
                f"Volatility: {'Yes' if last_session.get('volatility', False) else 'No'}",
                f"Quick Mode: {'Yes' if last_session.get('preview', False) else 'No'}"
            ])),
            ("Selected Plugins:", f"{len(last_session.get('selected_plugins', []))} plugins" if last_session.get('volatility', False) else "N/A"),
            ("Output File:", self._format_path(last_session.get('output_file', 'N/A')))
        ]
        
        for label, value in details:
            frame = ctk.CTkFrame(scroll_frame, fg_color="transparent")
            frame.pack(fill="x", pady=5)
            
            ctk.CTkLabel(
                frame,
                text=label,
                font=("Segoe UI", 12, "bold"),
                width=150,
                anchor="e"
            ).pack(side="left", padx=5)
            
            ctk.CTkLabel(
                frame,
                text=value,
                font=("Segoe UI", 12),
                wraplength=400,
                justify="left"
            ).pack(side="left", padx=5, fill="x", expand=True)
        
        # Close button
        ctk.CTkButton(
            container,
            text="Close",
            command=dialog.destroy
        ).pack(pady=10)
    
    def add_utility_buttons(self):
        self.sys_info_btn = ctk.CTkButton(
            self.main_container, 
            text="ℹ",
            font=("Segoe UI", 16, "bold"),
            fg_color=PRIMARY_COLOR, 
            hover_color=ACCENT_COLOR,
            width=40,
            height=40,
            corner_radius=20,
            command=self.show_system_info
        )
        self.sys_info_btn.place(relx=0.05, rely=0.05, anchor="nw")
        
        self.mode_btn = ctk.CTkButton(
            self.main_container, 
            text="☀" if ctk.get_appearance_mode() == "Dark" else "🌙",
            font=("Segoe UI", 16, "bold"),
            fg_color=PRIMARY_COLOR, 
            hover_color=ACCENT_COLOR,
            width=40,
            height=40,
            corner_radius=20,
            command=self.toggle_mode
        )
        self.mode_btn.place(relx=0.95, rely=0.05, anchor="ne")
    
    def show_system_info(self):
        ram = psutil.virtual_memory()
        cpu_usage = psutil.cpu_percent(interval=1)
        disk_usage = psutil.disk_usage('/')
        uptime = time.time() - psutil.boot_time()
        uptime_hours = int(uptime // 3600)
        uptime_minutes = int((uptime % 3600) // 60)
        
        info = (
            f"Total RAM: {ram.total / (1024 ** 3):.2f} GB\n"
            f"Used RAM: {ram.used / (1024 ** 3):.2f} GB\n"
            f"CPU Usage: {cpu_usage}%\n"
            f"Disk Usage: {disk_usage.percent}%\n"
            f"System Uptime: {uptime_hours}h {uptime_minutes}m"
        )
        show_info(info)
    
    def toggle_mode(self):
        global BACKGROUND_COLOR, PRIMARY_COLOR, SECONDARY_COLOR, ACCENT_COLOR, TEXT_COLOR
        current_mode = ctk.get_appearance_mode()
        if current_mode == "Dark":
            new_mode = "Light"
            BACKGROUND_COLOR = LIGHT_BACKGROUND
            PRIMARY_COLOR = LIGHT_PRIMARY
            SECONDARY_COLOR = LIGHT_SECONDARY
            ACCENT_COLOR = LIGHT_ACCENT
            TEXT_COLOR = LIGHT_TEXT
            self.mode_btn.configure(text="🌙")
        else:
            new_mode = "Dark"
            BACKGROUND_COLOR = DARK_BACKGROUND
            PRIMARY_COLOR = DARK_PRIMARY
            SECONDARY_COLOR = DARK_SECONDARY
            ACCENT_COLOR = DARK_ACCENT
            TEXT_COLOR = DARK_TEXT
            self.mode_btn.configure(text="☀")
        
        ctk.set_appearance_mode(new_mode)
        self.update_colors()
        
        self.config['appearance_mode'] = new_mode
        save_config(self.config)

    def update_colors(self):
        self.main_container.configure(fg_color=SECONDARY_COLOR)
        self.sys_info_btn.configure(fg_color=PRIMARY_COLOR, hover_color=ACCENT_COLOR)
        self.mode_btn.configure(fg_color=PRIMARY_COLOR, hover_color=ACCENT_COLOR)
        
        if hasattr(self, 'current_step'):
            if self.current_step == 0:
                self.update_welcome_colors()
            elif self.current_step == 1:
                self.update_step1_colors()
            elif self.current_step == 2:
                self.update_step2_colors()
            elif self.current_step == 3:
                self.update_step3_colors()
    
    def update_welcome_colors(self):
        if hasattr(self, 'welcome_label'):
            self.welcome_label.configure(text_color=TEXT_COLOR)
        if hasattr(self, 'line1'):
            self.line1.configure(text_color=PRIMARY_COLOR)
        if hasattr(self, 'line2'):
            self.line2.configure(text_color=PRIMARY_COLOR)
        if hasattr(self, 'start_btn'):
            self.start_btn.configure(fg_color=PRIMARY_COLOR, hover_color=ACCENT_COLOR)
        if hasattr(self, 'thanks_label'):
            self.thanks_label.configure(text_color=TEXT_COLOR)
    
    def update_step1_colors(self):
        if hasattr(self, 'investigator_name_entry'):
            entries = [
                getattr(self, 'investigator_name_entry', None),
                getattr(self, 'investigator_id_entry', None),
                getattr(self, 'case_name_entry', None),
                getattr(self, 'case_id_entry', None),
                getattr(self, 'memory_id_entry', None)
            ]
            
            for entry in entries:
                if entry:
                    try:
                        entry.configure(fg_color=SECONDARY_COLOR, text_color=TEXT_COLOR)
                    except:
                        pass
            
            if hasattr(self, 'report_format_menu'):
                try:
                    self.report_format_menu.configure(fg_color=SECONDARY_COLOR, text_color=TEXT_COLOR)
                except:
                    pass
            if hasattr(self, 'hash_menu'):
                try:
                    self.hash_menu.configure(fg_color=SECONDARY_COLOR, text_color=TEXT_COLOR)
                except:
                    pass
            
            if hasattr(self, 'input_frame'):
                try:
                    self.input_frame.configure(fg_color=BACKGROUND_COLOR)
                except:
                    pass
    
    def update_step2_colors(self):
        if hasattr(self, 'info_frame'):
            try:
                self.info_frame.configure(fg_color=BACKGROUND_COLOR)
                for label in self.info_frame.winfo_children():
                    if isinstance(label, ctk.CTkLabel):
                        label.configure(text_color=TEXT_COLOR)
            except:
                pass
        
        if hasattr(self, 'selected_file_label'):
            try:
                self.selected_file_label.configure(text_color=TEXT_COLOR)
            except:
                pass
        
        if hasattr(self, 'options_frame'):
            try:
                self.options_frame.configure(fg_color="transparent")
                for widget in self.options_frame.winfo_children():
                    if isinstance(widget, ctk.CTkCheckBox):
                        widget.configure(text_color=TEXT_COLOR)
            except:
                pass
    
    def update_step3_colors(self):
        if hasattr(self, 'info_frame'):
            try:
                self.info_frame.configure(fg_color=BACKGROUND_COLOR)
                for label in self.info_frame.winfo_children():
                    if isinstance(label, ctk.CTkLabel):
                        label.configure(text_color=TEXT_COLOR)
            except:
                pass
        
        if hasattr(self, 'progress_label'):
            try:
                self.progress_label.configure(text_color=TEXT_COLOR)
            except:
                pass
        
        if hasattr(self, 'time_left_label'):
            try:
                self.time_left_label.configure(text_color=TEXT_COLOR)
            except:
                pass
    
    def create_label_with_icon(self, parent, text, icon_key):
        frame = ctk.CTkFrame(parent, fg_color="transparent")
        
        if icon_key in self.icons:
            icon_label = ctk.CTkLabel(frame, image=self.icons[icon_key], text="")
            icon_label.pack(side="left", padx=(0, 5))
        
        label = ctk.CTkLabel(frame, text=text, font=("Segoe UI", 14), text_color=TEXT_COLOR)
        label.pack(side="left")
        
        return frame
    
    def show_step1(self):
        self.current_step = 1
        self.clear_container()
        
        container = ctk.CTkFrame(self.main_container, fg_color="transparent")
        container.pack(expand=True, fill="both", padx=15, pady=(45,15))
        
        title = ctk.CTkLabel(
            container,
            text="Step 1: Enter Case Details",
            font=("Segoe UI", 20, "bold"),
            text_color=PRIMARY_COLOR
        )
        title.pack(pady=(0,5))
        
        # Compact recent cases dropdown
        if self.config.get('recent_cases'):
            recent_frame = ctk.CTkFrame(container, fg_color="transparent")
            recent_frame.pack(fill="x", pady=(0,5))        
            self.create_label_with_icon(recent_frame, "Recent Cases:", "recent").pack(side="left", padx=(0, 2))
            
            self.recent_cases_var = StringVar(value="Select...")
            recent_menu = ctk.CTkComboBox(
                recent_frame,
                variable=self.recent_cases_var,
                values=[case['case_name'] for case in self.config['recent_cases']],
                command=self.load_recent_case,
                width=200,
                height=35,
                font=("Segoe UI", 14),
                dropdown_font=("Segoe UI", 12),
                fg_color=SECONDARY_COLOR,
                button_color=PRIMARY_COLOR,
                button_hover_color=ACCENT_COLOR,
            )
            recent_menu.pack(side="left", padx=(2,5), expand=True, fill="x")
        
        self.input_frame = ctk.CTkFrame(container, corner_radius=10, fg_color=BACKGROUND_COLOR)
        self.input_frame.pack(pady=5, padx=0, fill="both", expand=True)
        
        self.input_frame.grid_columnconfigure(0, weight=1)
        self.input_frame.grid_columnconfigure(1, weight=1)

        self.create_label_with_icon(self.input_frame, "Investigator Name:", "investigator_name").grid(row=0, column=0, padx=10, pady=8, sticky="e")
        self.investigator_name_entry = ctk.CTkEntry(self.input_frame, font=("Segoe UI", 14), fg_color=SECONDARY_COLOR, text_color=TEXT_COLOR)
        self.investigator_name_entry.grid(row=0, column=1, padx=10, pady=8, sticky="w")
        if 'investigator_name' in self.user_data:
            self.investigator_name_entry.insert(0, self.user_data['investigator_name'])
        self.investigator_name_entry.focus_set()

        self.create_label_with_icon(self.input_frame, "Investigator ID:", "investigator_id").grid(row=1, column=0, padx=10, pady=8, sticky="e")
        self.investigator_id_entry = ctk.CTkEntry(self.input_frame, font=("Segoe UI", 14), fg_color=SECONDARY_COLOR, text_color=TEXT_COLOR)
        self.investigator_id_entry.grid(row=1, column=1, padx=10, pady=8, sticky="w")
        if 'investigator_id' in self.user_data:
            self.investigator_id_entry.insert(0, self.user_data['investigator_id'])

        self.create_label_with_icon(self.input_frame, "Case Name:", "case_name").grid(row=2, column=0, padx=10, pady=8, sticky="e")
        self.case_name_entry = ctk.CTkEntry(self.input_frame, font=("Segoe UI", 14), fg_color=SECONDARY_COLOR, text_color=TEXT_COLOR)
        self.case_name_entry.grid(row=2, column=1, padx=10, pady=8, sticky="w")
        if 'case_name' in self.user_data:
            self.case_name_entry.insert(0, self.user_data['case_name'])

        self.create_label_with_icon(self.input_frame, "Case ID:", "case_id").grid(row=3, column=0, padx=10, pady=8, sticky="e")
        self.case_id_entry = ctk.CTkEntry(self.input_frame, font=("Segoe UI", 14), fg_color=SECONDARY_COLOR, text_color=TEXT_COLOR)
        self.case_id_entry.grid(row=3, column=1, padx=10, pady=8, sticky="w")
        if 'case_id' in self.user_data:
            self.case_id_entry.insert(0, self.user_data['case_id'])

        self.create_label_with_icon(self.input_frame, "Memory ID:", "memory_id").grid(row=4, column=0, padx=10, pady=8, sticky="e")
        self.memory_id_entry = ctk.CTkEntry(self.input_frame, font=("Segoe UI", 14), fg_color=SECONDARY_COLOR, text_color=TEXT_COLOR)
        self.memory_id_entry.grid(row=4, column=1, padx=10, pady=8, sticky="w")
        if 'memory_id' in self.user_data:
            self.memory_id_entry.insert(0, self.user_data['memory_id'])

        self.create_label_with_icon(self.input_frame, "Report Format:", "report_format").grid(row=5, column=0, padx=10, pady=8, sticky="e")
        self.report_format_var = StringVar(value="pdf")
        if 'report_format' in self.user_data:
            self.report_format_var.set(self.user_data['report_format'])
        self.report_format_menu = ctk.CTkOptionMenu(
            self.input_frame,
            variable=self.report_format_var,
            values=["txt", "pdf", "html", "docx", "json"],
            font=("Segoe UI", 14),
            fg_color=SECONDARY_COLOR,
            text_color=TEXT_COLOR
        )
        self.report_format_menu.grid(row=5, column=1, padx=10, pady=8, sticky="w")

        self.create_label_with_icon(self.input_frame, "Hash Algorithm:", "hash").grid(row=6, column=0, padx=10, pady=8, sticky="e")
        self.hash_algorithm = StringVar(value="sha256")
        if 'hash_algorithm' in self.user_data:
            self.hash_algorithm.set(self.user_data['hash_algorithm'])
        self.hash_menu = ctk.CTkOptionMenu(
            self.input_frame, 
            variable=self.hash_algorithm,
            values=["md5", "sha1", "sha256", "sha512"],
            font=("Segoe UI", 14),
            fg_color=SECONDARY_COLOR,
            text_color=TEXT_COLOR
        )
        self.hash_menu.grid(row=6, column=1, padx=10, pady=8, sticky="w")
        
        button_frame = ctk.CTkFrame(container, fg_color="transparent")
        button_frame.pack(pady=15)
        
        back_btn = ctk.CTkButton(
            button_frame,
            text="←",
            font=("Segoe UI", 18, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            width=60,
            height=60,
            corner_radius=30,
            command=self.show_welcome_screen
        )
        back_btn.pack(side="left", padx=10)
        
        next_btn = ctk.CTkButton(
            button_frame,
            text="→",
            font=("Segoe UI", 18, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            width=60,
            height=60,
            corner_radius=30,
            command=self.validate_step1
        )
        next_btn.pack(side="right", padx=10)
        
        self.investigator_name_entry.bind('<Return>', lambda e: self.investigator_id_entry.focus())
        self.investigator_id_entry.bind('<Return>', lambda e: self.case_name_entry.focus())
        self.case_name_entry.bind('<Return>', lambda e: self.case_id_entry.focus())
        self.case_id_entry.bind('<Return>', lambda e: self.memory_id_entry.focus())
        self.memory_id_entry.bind('<Return>', lambda e: self.report_format_menu.focus())
        self.report_format_menu.bind('<Return>', lambda e: self.hash_menu.focus())
        self.hash_menu.bind('<Return>', lambda e: self.validate_step1())
        
        self.add_utility_buttons()
    
    def load_recent_case(self, case_name):
        for case in self.config['recent_cases']:
            if case['case_name'] == case_name:
                self.user_data = case.copy()
                self.investigator_name_entry.delete(0, 'end')
                self.investigator_name_entry.insert(0, case['investigator_name'])
                self.investigator_id_entry.delete(0, 'end')
                self.investigator_id_entry.insert(0, case['investigator_id'])
                self.case_name_entry.delete(0, 'end')
                self.case_name_entry.insert(0, case['case_name'])
                self.case_id_entry.delete(0, 'end')
                self.case_id_entry.insert(0, case['case_id'])
                self.memory_id_entry.delete(0, 'end')
                self.memory_id_entry.insert(0, case['memory_id'])
                self.report_format_var.set(case.get('report_format', 'pdf'))
                self.hash_algorithm.set(case.get('hash_algorithm', 'sha256'))
                break
    
    def validate_step1(self):
        if not all([
            self.investigator_name_entry.get(),
            self.investigator_id_entry.get(),
            self.case_name_entry.get(),
            self.case_id_entry.get(),
            self.memory_id_entry.get()
        ]):
            show_error("Please fill all fields")
            return
        
        self.user_data = {
            "investigator_name": self.investigator_name_entry.get(),
            "investigator_id": self.investigator_id_entry.get(),
            "case_name": self.case_name_entry.get(),
            "case_id": self.case_id_entry.get(),
            "memory_id": self.memory_id_entry.get(),
            "report_format": self.report_format_var.get(),
            "hash_algorithm": self.hash_algorithm.get()
        }
        
        # Add to recent cases (limit to MAX_RECENT_CASES)
        if not any(case['case_name'] == self.user_data['case_name'] for case in self.config['recent_cases']):
            self.config['recent_cases'].insert(0, self.user_data.copy())
            self.config['recent_cases'] = self.config['recent_cases'][:MAX_RECENT_CASES]
            save_config(self.config)
        
        self.show_step2()
    
    def show_step2(self):
        self.current_step = 2
        self.clear_container()
        
        container = ctk.CTkFrame(self.main_container, fg_color="transparent")
        container.pack(expand=True, fill="both", padx=20, pady=20)
        
        title = ctk.CTkLabel(
            container,
            text="Step 2: Select Output and Options",
            font=("Segoe UI", 20, "bold"),
            text_color=PRIMARY_COLOR
        )
        title.pack(pady=20)
        
        # Get memory information
        mem = psutil.virtual_memory()
        ram_size_gb = mem.total / (1024 ** 3)
        
        # Calculate estimated dump size (20% overhead)
        estimated_dump_size = mem.total * 1.2
        estimated_dump_size_gb = estimated_dump_size / (1024 ** 3)
        
        # Check disk space if output file is selected
        disk_warning = ""
        if self.output_file:
            output_dir = os.path.dirname(self.output_file)
            try:
                disk_free = psutil.disk_usage(output_dir).free
                disk_free_gb = disk_free / (1024 ** 3)
                if disk_free < estimated_dump_size:
                    disk_warning = (
                        f"\n\n🚨 CRITICAL: Only {disk_free_gb:.1f} GB free in output directory!"
                        f"\nYou need at least {estimated_dump_size_gb:.1f} GB of free space."
                    )
            except Exception as e:
                logging.error(f"Could not check disk space: {e}")
        
        # Build warning message
        warning_text = (
            f"⚠ System RAM: {ram_size_gb:.1f} GB"
            f" ⚠ Estimated dump size: {estimated_dump_size_gb:.1f} GB"
        )
        
        # Add compression estimate if enabled
        if self.enable_compression:
            compressed_estimate = estimated_dump_size * 0.7  # 30% compression
            warning_text += f"\n\nWith compression: ~{compressed_estimate/(1024**3):.1f} GB"
        
        # Add recommendations for large systems
        if ram_size_gb > 8:
            warning_text += "\nConsider enabling compression to reduce file size."
            if ram_size_gb > 16:
                warning_text += "\nQuick mode recommended for faster capture."
        
        # Add disk space warning if needed
        warning_text += disk_warning
        
        warning_label = ctk.CTkLabel(
            container,
            text=warning_text,
            font=("Segoe UI", 12),
            text_color="#FFA500",  # Orange color for warning
            justify="left"
        )
        warning_label.pack(pady=5)
        
        self.info_frame = ctk.CTkFrame(container, corner_radius=10, fg_color=BACKGROUND_COLOR)
        self.info_frame.pack(pady=10, padx=20, fill="x")
        
        ctk.CTkLabel(
            self.info_frame, 
            text=f"Investigator: {self.user_data['investigator_name']} (ID: {self.user_data['investigator_id']})",
            font=("Segoe UI", 14),
            text_color=TEXT_COLOR
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.info_frame, 
            text=f"Case: {self.user_data['case_name']} (ID: {self.user_data['case_id']})",
            font=("Segoe UI", 14),
            text_color=TEXT_COLOR
        ).pack(pady=5)
        
        select_btn_frame = ctk.CTkFrame(container, fg_color="transparent")
        select_btn_frame.pack(pady=20)
        
        select_btn = ctk.CTkButton(
            select_btn_frame,
            text="📁 Select Output File",
            font=("Segoe UI", 14, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            command=self.select_file
        )
        select_btn.pack(pady=10)
        
        self.selected_file_label = ctk.CTkLabel(
            select_btn_frame,
            text="No file selected" if not self.output_file else self._format_path(self.output_file),
            font=("Segoe UI", 12),
            text_color=TEXT_COLOR,
            justify="left",
            wraplength=400
        )
        self.selected_file_label.pack(pady=5)
        
        self.options_frame = ctk.CTkFrame(container, fg_color="transparent")
        self.options_frame.pack(pady=10, fill="x")
        
        centered_frame = ctk.CTkFrame(self.options_frame, fg_color="transparent")
        centered_frame.pack(pady=5)
        
        self.compression_var = ctk.BooleanVar(value=self.enable_compression)
        self.enable_compression_check = ctk.CTkCheckBox(
            centered_frame,
            text="Compress RAM image after capture",
            variable=self.compression_var,
            command=self.toggle_compression
        )
        self.enable_compression_check.pack(pady=5)
        
        self.volatility_var = ctk.BooleanVar(value=self.enable_volatility)
        self.enable_volatility_check = ctk.CTkCheckBox(
            centered_frame,
            text="Analyze with Volatility",
            variable=self.volatility_var,
            command=self.toggle_volatility
        )
        self.enable_volatility_check.pack(pady=5)
        
        # Plugin selection frame
        self.plugin_frame = ctk.CTkFrame(self.options_frame, fg_color="transparent")
        
        def update_plugin_frame_visibility():
            if self.volatility_var.get():
                self.plugin_frame.pack(pady=5, fill="x")
            else:
                self.plugin_frame.pack_forget()
        
        self.volatility_var.trace_add('write', lambda *_: update_plugin_frame_visibility())
        
        # Plugin selection button
        self.plugin_btn = ctk.CTkButton(
            self.plugin_frame,
            text="⚙ Select Volatility Plugins",
            font=("Segoe UI", 12),
            fg_color=SECONDARY_COLOR,
            hover_color=ACCENT_COLOR,
            command=self.show_plugin_selection
        )
        self.plugin_btn.pack(pady=5)
        
        # Show currently selected plugins count
        self.plugin_count_label = ctk.CTkLabel(
            self.plugin_frame,
            text=f"Selected: {len(self.selected_plugins)} plugins",
            font=("Segoe UI", 10),
            text_color=TEXT_COLOR
        )
        self.plugin_count_label.pack(pady=2)
        
        # Update visibility initially
        update_plugin_frame_visibility()
        
        self.preview_var = ctk.BooleanVar(value=self.enable_preview)
        self.enable_preview_check = ctk.CTkCheckBox(
            centered_frame,
            text="Quick Capture Mode",
            variable=self.preview_var,
            command=self.toggle_preview
        )
        self.enable_preview_check.pack(pady=5)
        
        button_frame = ctk.CTkFrame(container, fg_color="transparent")
        button_frame.pack(pady=20)
        
        back_btn = ctk.CTkButton(
            button_frame,
            text="←",
            font=("Segoe UI", 18, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            width=60,
            height=60,
            corner_radius=30,
            command=self.show_step1
        )
        back_btn.pack(side="left", padx=10)
        
        self.next_btn_step2 = ctk.CTkButton(
            button_frame,
            text="→",
            font=("Segoe UI", 18, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            width=60,
            height=60,
            corner_radius=30,
            state="disabled" if not self.output_file else "normal",
            command=self.show_step3
        )
        self.next_btn_step2.pack(side="right", padx=10)
        
        self.root.bind('<Return>', lambda event: self.select_file())
        self.add_utility_buttons()
    
    def show_plugin_selection(self):
        """Show dialog for selecting Volatility plugins"""
        dialog = ctk.CTkToplevel(self.root)
        dialog.title("Select Volatility Plugins")
        dialog.geometry("600x600")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Main container
        container = ctk.CTkFrame(dialog)
        container.pack(expand=True, fill="both", padx=10, pady=10)
        
        # Select all button
        select_all_frame = ctk.CTkFrame(container, fg_color="transparent")
        select_all_frame.pack(fill="x", pady=5)
        
        ctk.CTkButton(
            select_all_frame,
            text="Select All",
            width=100,
            command=lambda: [var.set(True) for var in self.plugin_vars.values()]
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            select_all_frame,
            text="Deselect All",
            width=100,
            command=lambda: [var.set(False) for var in self.plugin_vars.values()]
        ).pack(side="left", padx=5)
        
        # Plugin list in scrollable frame
        scroll_frame = ctk.CTkScrollableFrame(container)
        scroll_frame.pack(expand=True, fill="both", padx=5, pady=5)
        
        self.plugin_vars = {}
        for plugin, details in self.all_plugins.items():
            plugin_name = plugin.split('.')[-1]
            var = ctk.BooleanVar(value=plugin in self.selected_plugins)
            self.plugin_vars[plugin] = var
            
            # Add note for slow plugins
            slow_note = " (may take longer)" if details['slow'] else ""
            
            cb = ctk.CTkCheckBox(
                scroll_frame,
                text=f"{plugin_name}{slow_note}",
                variable=var
            )
            cb.pack(anchor="w", pady=2)
        
        # Save button
        def save_selection():
            self.selected_plugins = [p for p, var in self.plugin_vars.items() if var.get()]
            self.plugin_count_label.configure(text=f"Selected: {len(self.selected_plugins)} plugins")
            
            # Save to config
            self.config['volatility_plugins'] = self.selected_plugins
            save_config(self.config)
            dialog.destroy()
        
        ctk.CTkButton(
            container,
            text="Save Selection",
            command=save_selection
        ).pack(pady=10)
    
    def run_live_analysis(self):
        """Perform live process analysis"""
        try:
            self.live_analysis_active = True
            output_dir = os.path.join(BASE_DIR, "LiveAnalysis")
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            output_file = os.path.join(output_dir, f"live_processes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(f"Live Process Analysis - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("="*80 + "\n\n")
                
                f.write("Running Processes:\n")
                f.write("="*40 + "\n")
                for proc in psutil.process_iter(['pid', 'name', 'username', 'create_time']):
                    try:
                        create_time = datetime.fromtimestamp(proc.create_time()).strftime('%Y-%m-%d %H:%M:%S')
                        f.write(f"{proc.pid:>6} {proc.info['name']:20} {proc.info['username']:20} {create_time}\n")
                    except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                        pass
                
                f.write("\nNetwork Connections:\n")
                f.write("="*40 + "\n")
                for conn in psutil.net_connections(kind='inet'):
                    if conn.status == psutil.CONN_ESTABLISHED:
                        f.write(f"{conn.laddr.ip}:{conn.laddr.port} -> {conn.raddr.ip}:{conn.raddr.port} ({conn.pid})\n")
            
            show_info(f"Live analysis saved to:\n{output_file}")
            os.startfile(output_file)
            
        except Exception as e:
            show_error(f"Live analysis failed: {str(e)}")
        finally:
            self.live_analysis_active = False
    
    def toggle_compression(self):
        self.enable_compression = self.compression_var.get()
        logging.info(f"Compression {'enabled' if self.enable_compression else 'disabled'}")
    
    def toggle_volatility(self):
        if self.volatility_var.get():  # If user is enabling Volatility
            if not check_python_installed():
                show_error("Python is required for Volatility but was not found.\n"
                         "Please install Python and ensure it's in your PATH.")
                self.volatility_var.set(False)  # Uncheck the option
                return
        self.enable_volatility = self.volatility_var.get()
        logging.info(f"Volatility analysis {'enabled' if self.enable_volatility else 'disabled'}")
    
    def toggle_preview(self):
        self.enable_preview = self.preview_var.get()
        logging.info(f"Preview mode {'enabled' if self.enable_preview else 'disabled'}")
    
    def select_file(self):
        # Generate default filename
        default_name = generate_filename(
            self.user_data["case_name"],
            self.user_data["case_id"],
            self.user_data["memory_id"]
        )
        
        file_path = filedialog.asksaveasfilename(
            title="Select location to save memory image",
            defaultextension=".raw",
            initialfile=default_name,
            filetypes=[("RAW files", "*.raw"), ("All Files", "*.*")]
        )
        
        if file_path:
            self.output_file = file_path
            self.selected_file_label.configure(text=f"Selected:\n{self._format_path(self.output_file)}")
            self.next_btn_step2.configure(state="normal")
            self.root.unbind('<Return>')
            self.root.bind('<Return>', lambda event: self.show_step3())
        else:
            show_error("No file selected. Please try again.")
    
    def get_volatility_command(self, image_path, plugin):
        """Helper method to get volatility command for a plugin"""
        # Create required directories if they don't exist
        shutil.rmtree(CACHE_DIR, ignore_errors=True)
        CACHE_DIR.mkdir(exist_ok=True)
        
        if (VOLATILITY_PATH / "vol.exe").exists():
            base_cmd = [str(VOLATILITY_PATH / "vol.exe")]
        elif (VOLATILITY_PATH / "vol.py").exists():
            base_cmd = ['python', str(VOLATILITY_PATH / 'vol.py')]
        else:
            return None
        
        # Common parameters for all commands
        cmd = base_cmd + [
            '-f', image_path,
            plugin
        ]
        
        
        
        return cmd

    
    def analyze_with_volatility(self, image_path, output_dir):
        """Analyze memory dump with selected Volatility plugins"""
        try:
            # First try to run windows.info to initialize symbols
            info_cmd = self.get_volatility_command(image_path, "windows.info.Info")
            if info_cmd:
                info_result = subprocess.run(
                    info_cmd,
                    capture_output=True,
                    text=True,
                    encoding='utf-8',
                    errors='replace',
                    timeout=60
                )
                if "Unsatisfied requirement" in info_result.stderr:
                    show_error("Volatility symbol initialization failed. Trying to continue...")
            
            if not self.selected_plugins:
                show_error("No Volatility plugins selected for analysis")
                return None
            
            output_file = os.path.join(output_dir, f"volatility_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(f"Volatility Analysis Report - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Memory Image: {image_path}\n")
                f.write("="*80 + "\n\n")
                
                total_plugins = len(self.selected_plugins)
                
                for i, plugin in enumerate(self.selected_plugins):
                    plugin_name = plugin.split('.')[-1]
                    start_time = time.time()
                    
                    # Update progress
                    progress = i / total_plugins
                    self.root.after(0, lambda: self.finalizing_progress_bar.set(progress))
                    
                    # Show plugin info
                    status_text = f"Running {plugin_name}... ({i+1}/{total_plugins})"
                    self.root.after(0, lambda: self.finalizing_time_label.configure(text=status_text))
                    
                    # Run the plugin
                    cmd = self.get_volatility_command(image_path, plugin)
                    if not cmd:
                        continue
                    
                    result = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        encoding='utf-8',
                        errors='replace',
                        timeout=300  # 5 minute timeout per plugin
                    )
                    
                    # Write results to report
                    f.write(f"\n=== {plugin_name} ===\n")
                    f.write(result.stdout)
                    if result.stderr:
                        f.write(f"\nErrors:\n{result.stderr}\n")
                    
                    # Update progress
                    progress = (i + 1) / total_plugins
                    self.root.after(0, lambda: self.finalizing_progress_bar.set(progress))
                    
                    elapsed = time.time() - start_time
                    logging.info(f"Plugin {plugin_name} completed in {elapsed:.1f}s")
            
            logging.info(f"Volatility analysis saved to {output_file}")
            return output_file
        except Exception as e:
            logging.error(f"Volatility analysis failed: {e}")
            return None

    def clear_volatility_cache(self):
        """Clear Volatility cache to force fresh symbol downloads"""
        try:
            if CACHE_DIR.exists():
                shutil.rmtree(CACHE_DIR)
            CACHE_DIR.mkdir()
            return True
        except Exception as e:
            logging.error(f"Error clearing cache: {e}")
            return False
       
    
    def analyze_process_timeline(self, image_path):
        """Generate and display process timeline"""
        try:
            output_dir = os.path.dirname(image_path)
            timeline_file = os.path.join(output_dir, f"process_timeline_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html")
            
            # Get process list with start times - fixed command

            # Run info plugin first to initialize context
            info_cmd = self.get_volatility_command(image_path, "windows.info.Info")
            subprocess.run(info_cmd, capture_output=True, text=True)

            # Now run pslist
            cmd = self.get_volatility_command(image_path, "windows.pslist.PsList")

            if not cmd:
                return False
                
            result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
            
            if result.returncode != 0:
                show_error(f"Failed to get process list:\n{result.stderr}")
                return False
            
            # Parse process information
            processes = []
            for line in result.stdout.split('\n'):
                if not line.strip() or 'PID' in line or '---' in line:
                    continue
                parts = [p for p in line.split(' ') if p]
                if len(parts) >= 6:
                    try:
                        pid = parts[0]
                        ppid = parts[1]
                        start_time = ' '.join(parts[2:5])
                        name = parts[5]
                        processes.append({
                            'pid': pid,
                            'ppid': ppid,
                            'start_time': start_time,
                            'name': name
                        })
                    except:
                        continue
            
            # Generate HTML timeline
            html_content = f"""<!DOCTYPE html>
            <html>
            <head>
                <title>Process Timeline</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; }}
                    .timeline {{ position: relative; max-width: 1200px; margin: 0 auto; }}
                    .timeline::after {{ content: ''; position: absolute; width: 6px; background-color: #3498db; top: 0; bottom: 0; left: 50%; margin-left: -3px; }}
                    .container {{ padding: 10px 40px; position: relative; background-color: inherit; width: 50%; }}
                    .container::after {{ content: ''; position: absolute; width: 25px; height: 25px; right: -17px; background-color: white; border: 4px solid #FF9F55; top: 15px; border-radius: 50%; z-index: 1; }}
                    .left {{ left: 0; }}
                    .right {{ left: 50%; }}
                    .left::before {{ content: " "; height: 0; position: absolute; top: 22px; width: 0; z-index: 1; right: 30px; border: medium solid #f1f1f1; border-width: 10px 0 10px 10px; border-color: transparent transparent transparent #f1f1f1; }}
                    .right::before {{ content: " "; height: 0; position: absolute; top: 22px; width: 0; z-index: 1; left: 30px; border: medium solid #f1f1f1; border-width: 10px 10px 10px 0; border-color: transparent #f1f1f1 transparent transparent; }}
                    .content {{ padding: 20px 30px; background-color: white; position: relative; border-radius: 6px; box-shadow: 0 2px 5px rgba(0,0,0,0.1); }}
                    h2 {{ color: #3498db; }}
                    .process-info {{ margin-top: 5px; font-size: 14px; color: #555; }}
                </style>
            </head>
            <body>
                <div class="timeline">
                    <h1>Process Timeline</h1>
                    <p>Memory Image: {image_path}</p>
                    <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            """
            
            for i, process in enumerate(processes):
                side = "left" if i % 2 == 0 else "right"
                html_content += f"""
                    <div class="container {side}">
                        <div class="content">
                            <h2>{process['name']}</h2>
                            <div class="process-info">
                                <p><strong>Started:</strong> {process['start_time']}</p>
                                <p><strong>PID:</strong> {process['pid']} | <strong>PPID:</strong> {process['ppid']}</p>
                            </div>
                        </div>
                    </div>
                """
            
            html_content += """
                </div>
            </body>
            </html>
            """
            
            with open(timeline_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            show_info(f"Process timeline generated:\n{timeline_file}")
            os.startfile(timeline_file)
            return True
            
        except Exception as e:
            logging.error(f"Timeline error: {e}")
            show_error(f"Error generating timeline: {str(e)}")
            return False
    
    def check_for_anti_forensics(self, image_path):
        """Check for signs of anti-forensics activity"""
        try:
            output_dir = os.path.dirname(image_path)
            report_file = os.path.join(output_dir, f"anti_forensics_check_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
            
            suspicious_items = []
            
            # Check for processes with no parent
            cmd = self.get_volatility_command(image_path, "windows.pstree.PsTree")
            if cmd:
                result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
                if result.returncode == 0:
                    for line in result.stdout.split('\n'):
                        if "-> ???" in line:
                            suspicious_items.append(f"Orphan process: {line.strip()}")
            
            # Check for suspicious DLLs
            cmd = self.get_volatility_command(image_path, "windows.dlllist.DllList")
            if cmd:
                result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
                if result.returncode == 0:
                    for line in result.stdout.split('\n'):
                        if any(x in line.lower() for x in ['inject', 'hook', 'rootkit']):
                            suspicious_items.append(f"Suspicious DLL: {line.strip()}")
            
            # Save report
            with open(report_file, 'w', encoding='utf-8') as f:
                f.write("Anti-Forensics Check Report\n")
                f.write("="*40 + "\n")
                f.write(f"Memory Image: {image_path}\n")
                f.write(f"Analysis Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                if suspicious_items:
                    f.write("⚠️ SUSPICIOUS ITEMS FOUND:\n")
                    f.write("="*40 + "\n")
                    for item in suspicious_items:
                        f.write(f"- {item}\n")
                else:
                    f.write("✅ No obvious signs of anti-forensics activity detected\n")
            
            if suspicious_items:
                show_info(f"Anti-forensics check completed with findings:\n{report_file}")
            else:
                show_info("No signs of anti-forensics activity detected")
            
            os.startfile(report_file)
            return True
            
        except Exception as e:
            show_error(f"Error during anti-forensics check: {str(e)}")
            return False
    
    def show_step3(self):
        self.current_step = 3
        self.clear_container()
        
        container = ctk.CTkFrame(self.main_container, fg_color="transparent")
        container.pack(expand=True, fill="both", padx=20, pady=20)
        
        title = ctk.CTkLabel(
            container,
            text="Step 3: Capture Memory",
            font=("Segoe UI", 20, "bold"),
            text_color=PRIMARY_COLOR
        )
        title.pack(pady=20)
        
        self.info_frame = ctk.CTkFrame(container, corner_radius=10, fg_color=BACKGROUND_COLOR)
        self.info_frame.pack(pady=10, padx=20, fill="x")
        
        ctk.CTkLabel(
            self.info_frame, 
            text=f"Case: {self.user_data['case_name']}",
            font=("Segoe UI", 14),
            text_color=TEXT_COLOR
        ).pack(pady=5)
        
        ctk.CTkLabel(
            self.info_frame, 
            text=f"Output File:\n{self._format_path(self.output_file)}",
            font=("Segoe UI", 14),
            text_color=TEXT_COLOR,
            justify="left"
        ).pack(pady=5)
        
        features = []
        if self.enable_compression:
            features.append("Compression")
        if self.enable_volatility:
            features.append("Volatility Analysis")
        if self.enable_preview:
            features.append("Quick Capture Mode")
        
        if features:
            ctk.CTkLabel(
                self.info_frame,
                text=f"Enabled features: {', '.join(features)}",
                font=("Segoe UI", 12),
                text_color=TEXT_COLOR
            ).pack(pady=5)
        
        # Add Live Analysis button
        live_btn = ctk.CTkButton(
            container,
            text="🔍",
            font=("Segoe UI", 18, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            width=60,
            height=60,
            corner_radius=30,
            command=self.run_live_analysis
        )
        live_btn.pack(pady=10)
        
        progress_container = ctk.CTkFrame(container, fg_color="transparent")
        progress_container.pack(pady=40, fill="x", expand=True)
        
        self.progress_bar = ctk.CTkProgressBar(
            progress_container,
            mode="determinate",
            width=500,
            height=30,
            fg_color=SECONDARY_COLOR,
            progress_color=PRIMARY_COLOR
        )
        self.progress_bar.pack(pady=20)
        self.progress_bar.set(0)
        
        self.progress_label = ctk.CTkLabel(
            progress_container,
            text="Ready to capture... 0%",
            font=("Segoe UI", 16),
            text_color=TEXT_COLOR
        )
        self.progress_label.pack(pady=10)
        
        self.time_left_label = ctk.CTkLabel(
            progress_container,
            text="Estimated time remaining: --",
            font=("Segoe UI", 14),
            text_color=TEXT_COLOR
        )
        self.time_left_label.pack(pady=5)
        
        button_frame = ctk.CTkFrame(container, fg_color="transparent")
        button_frame.pack(side="bottom", pady=20)
        
        back_btn = ctk.CTkButton(
            button_frame,
            text="←",
            font=("Segoe UI", 18, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            width=60,
            height=60,
            corner_radius=30,
            command=self.show_step2
        )
        back_btn.pack(side="left", padx=10)
        
        self.capture_btn = ctk.CTkButton(
            button_frame,
            text="⚡",
            font=("Segoe UI", 18, "bold"),
            fg_color=PRIMARY_COLOR,
            hover_color=ACCENT_COLOR,
            width=60,
            height=60,
            corner_radius=30,
            command=self.start_capture
        )
        self.capture_btn.pack(side="right", padx=10)
        
        self.root.bind('<Return>', lambda event: self.start_capture())
        self.add_utility_buttons()

    def estimate_remaining_time(self, progress):
        if progress <= 0:
            return "--"
        
        if hasattr(self, 'finalizing_start_time'):
            elapsed = time.time() - self.finalizing_start_time
            remaining = (elapsed / progress) * (1 - progress)
            
            if remaining > 60:
                mins = int(remaining // 60)
                secs = int(remaining % 60)
                return f"{mins}m {secs}s"
            else:
                return f"{remaining:.1f}s"
        return "--"

    def start_capture(self):
        # Final checks before starting capture
        if self.enable_volatility and not check_python_installed():
            show_error("Python is required for Volatility but was not found.\n"
                     "Please install Python and ensure it's in your PATH.")
            return
        
        # Reset capture variables
        self.capture_active = True
        self.capture_bytes_copied = 0
        self.capture_start_time = time.time()
        self.total_memory_size = psutil.virtual_memory().total
        
        # Update UI
        self.capture_btn.configure(text="■", command=self.stop_capture)
        self.progress_label.configure(text="Initializing capture... 0%")
        self.progress_bar.set(0)
        self.time_left_label.configure(text="Estimated time remaining: --")
        
        # Build the command
        command = [str(BASE_DIR / "winpmem_mini.exe"), self.output_file]
        if self.enable_preview:
            command.append("--preview")
        
        # Start the capture thread
        self.capture_thread = threading.Thread(
            target=self.run_capture_process,
            args=(command,),
            daemon=True
        )
        self.capture_thread.start()

    def execute_next_finalizing_step(self, start_time):
        if self.current_finalizing_step < len(self.finalizing_steps):
            step_name, step_func = self.finalizing_steps[self.current_finalizing_step]
            self.finalizing_progress_label.configure(text=step_name)
            
            # Calculate total time estimate for remaining steps
            total_steps = len(self.finalizing_steps)
            current_step = self.current_finalizing_step
            remaining_steps = total_steps - current_step
            
            # Update progress bar
            progress = current_step / total_steps
            self.finalizing_progress_bar.set(progress)
            
            # Update ETA
            if step_name.startswith("Running Volatility"):
                # For volatility, we'll update ETA per plugin in the analyze_with_volatility method
                self.finalizing_time_label.configure(text="Starting analysis...")
            else:
                # For other steps, show generic message
                self.finalizing_time_label.configure(text=f"Starting {step_name}...")
            
            # Run the step in a thread
            threading.Thread(
                target=self.run_finalizing_step,
                args=(step_name, step_func, start_time),
                daemon=True
            ).start()
        else:
            # All steps completed
            total_time = time.time() - start_time
            self.finalize_processing(start_time, total_time)

    def run_finalizing_step(self, step_name, step_func, start_time):
        step_start = time.time()
        try:
            result = step_func()
            
            # Update progress
            self.current_finalizing_step += 1
            progress = self.current_finalizing_step / len(self.finalizing_steps)
            
            self.root.after(0, lambda: self.finalizing_progress_bar.set(progress))
            self.root.after(0, lambda: self.finalizing_time_label.configure(
                text=f"Completed {step_name} in {time.time() - step_start:.1f}s"
            ))
            
            # Execute next step
            self.root.after(0, lambda: self.execute_next_finalizing_step(start_time))
            
            return result
        except Exception as e:
            self.root.after(0, lambda: show_error(f"Error in {step_name}: {str(e)}"))
            self.root.after(0, lambda: self.finish_capture(stopped=True))

    def run_capture_process(self, command):
        start_time = time.time()
        last_update_time = start_time
        last_size = 0
        
        mem_info = psutil.virtual_memory()
        self.total_memory_size = mem_info.total
        self.capture_start_time = start_time
        
        try:
            startupinfo = None
            if os.name == 'nt':
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            self.capture_process = subprocess.Popen(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                startupinfo=startupinfo
            )
            
            time.sleep(1)
            
            while self.capture_active:
                current_time = time.time()
                
                if self.capture_process.poll() is not None:
                    break
                
                try:
                    current_size = os.path.getsize(self.output_file)
                except FileNotFoundError:
                    current_size = 0
                
                progress = min((current_size) / self.total_memory_size, 1.0)
                
                if current_time - last_update_time >= 0.5 or abs(current_size - last_size) > (self.total_memory_size * 0.01):
                    self.progress_label.configure(text=f"Capturing RAM... {progress*100:.1f}%")
                    self.progress_bar.set(progress)
                    
                    elapsed = current_time - start_time
                    if progress > 0.01 and elapsed > 0:
                        transfer_rate = (current_size - last_size) / (1024 * 1024 * (current_time - last_update_time)) if current_time > last_update_time else 0
                        remaining = max((self.total_memory_size - current_size) / (current_size / elapsed) if current_size > 0 else 0, 0)
                        
                        # Convert to human-readable time
                        if remaining > 60:
                            mins = int(remaining // 60)
                            secs = int(remaining % 60)
                            time_str = f"{mins}m {secs}s"
                        else:
                            time_str = f"{remaining:.1f}s"
                        
                        self.time_left_label.configure(text=f"ETA: {time_str} | Speed: {transfer_rate:.2f} MB/s")
                    
                    last_size = current_size
                    last_update_time = current_time
                
                time.sleep(0.1)
            
            if not self.capture_active:
                self.root.after(0, lambda: self.finish_capture(stopped=True))
                return
            
            # Hide capture progress bar and show finalizing progress
            self.progress_bar.pack_forget()
            self.progress_label.pack_forget()
            self.time_left_label.pack_forget()
            
            # Create finalizing progress UI
            self.finalizing_progress_bar = ctk.CTkProgressBar(
                self.progress_bar.master,
                mode="determinate",
                width=500,
                height=30,
                fg_color=SECONDARY_COLOR,
                progress_color=PRIMARY_COLOR
            )
            self.finalizing_progress_bar.pack(pady=20)
            self.finalizing_progress_bar.set(0)
            
            self.finalizing_progress_label = ctk.CTkLabel(
                self.progress_label.master,
                text="Finalizing...",
                font=("Segoe UI", 16),
                text_color=TEXT_COLOR
            )
            self.finalizing_progress_label.pack(pady=10)
            
            self.finalizing_time_label = ctk.CTkLabel(
                self.time_left_label.master,
                text="Preparing final steps...",
                font=("Segoe UI", 14),
                text_color=TEXT_COLOR
            )
            self.finalizing_time_label.pack(pady=5)
            
            self.finalizing_start_time = time.time()
            self.finalizing_steps = []
            
            # Define finalizing steps with weights for progress bar
            total_weight = 0
            if self.enable_volatility:
                total_weight += len(self.selected_plugins)
            if self.enable_compression:
                total_weight += 1
            
            current_weight = 0
            
            if self.enable_volatility:
                # Add volatility analysis as a step
                self.finalizing_steps.append((
                    "Running Volatility Analysis", 
                    lambda: self.analyze_with_volatility(self.output_file, os.path.dirname(self.output_file))
                ))
            
            if self.enable_compression:
                # Add compression as a step
                self.finalizing_steps.append((
                    "Compressing Files", 
                    self.run_compression
                ))
            
            self.current_finalizing_step = 0
            self.execute_next_finalizing_step(start_time)
            
        except Exception as e:
            self.root.after(0, lambda: self.progress_label.configure(text=f"Error: {str(e)}"))
            self.root.after(0, lambda: show_error(f"An exception occurred: {str(e)}"))
        finally:
            self.capture_active = False
            if hasattr(self, 'capture_process') and self.capture_process:
                try:
                    self.capture_process.terminate()
                except:
                    pass

    def run_volatility_analysis(self):
        if not hasattr(self, 'output_file'):
            return None

        output_dir = os.path.dirname(self.output_file)
        os.makedirs(output_dir, exist_ok=True)
        report_file = os.path.join(output_dir, f"volatility_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")

        start_time = time.time()
        plugin_times = []

        try:
            with open(report_file, 'w', encoding='utf-8') as report:
                total_plugins = len(self.selected_plugins)
                for i, plugin in enumerate(self.selected_plugins):
                    plugin_name = plugin.split('.')[-1]
                    plugin_start = time.time()

                    cmd = self.get_volatility_command(self.output_file, plugin)
                    if not cmd:
                        continue

                    process = subprocess.Popen(
                        cmd,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                        universal_newlines=True,
                        encoding='utf-8'
                    )

                    while True:
                        line = process.stdout.readline()
                        if not line and process.poll() is not None:
                            break

                        report.write(line)
                        report.flush()

                        if "Progress:" in line:
                            try:
                                percent = float(line.split("Progress:")[1].split("%")[0].strip())
                                step_progress = (i + percent / 100) / total_plugins
                                elapsed = time.time() - start_time
                                avg_time = elapsed / max((i + percent / 100), 0.01)
                                remaining = avg_time * (total_plugins - (i + percent / 100))
                                eta_str = f"{int(remaining//60)}m {int(remaining%60)}s" if remaining > 60 else f"{remaining:.1f}s"

                                self.root.after(0, lambda: self.finalizing_progress_bar.set(step_progress))
                                self.root.after(0, lambda: self.finalizing_time_label.configure(
                                    text=f"{plugin_name} ({percent:.1f}%) | ETA: {eta_str}"
                                ))
                            except:
                                pass

                    plugin_times.append(time.time() - plugin_start)

            self.analysis_time = time.time() - start_time
            self.volatility_analysis_time = time.time() - start_time
            self.volatility_time = time.time() - start_time  # <-- Store total time in seconds
            self.volatility_time_str = f"{int(self.volatility_time//60)}m {int(self.volatility_time%60)}s"  # <-- Formatted string
            return report_file
        except Exception as e:
            logging.error(f"Volatility analysis failed: {e}")
            return None

    def run_compression(self):
        """Handle compression with progress tracking"""
        try:
            start_time = time.time()
            original_file = self.output_file
            zip_path = original_file + ".zip"
            total_size = os.path.getsize(original_file)
            
            # Update UI
            self.root.after(0, lambda: [
                self.finalizing_progress_bar.set(0),
                self.finalizing_time_label.configure(
                    text="Starting compression (ZIP64)" if total_size > 4*1024*1024*1024 
                    else "Starting compression"
                )
            ])

            processed = 0
            chunk_size = 1024 * 1024 * 4  # 4MB chunks

            # Create new ZIP file
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED, allowZip64=True) as zipf:
                with open(original_file, 'rb') as f_in:
                    with zipf.open(os.path.basename(original_file), 'w', force_zip64=True) as f_out:
                        while True:
                            chunk = f_in.read(chunk_size)
                            if not chunk:
                                break
                            f_out.write(chunk)
                            processed += len(chunk)
                            progress = processed / total_size
                            self.update_compression_ui(progress)

            # Verify the compressed file
            if os.path.exists(zip_path) and os.path.getsize(zip_path) > 0:
                # Delete original file only after successful compression
                os.remove(original_file)
                # Update output_file reference to point to the compressed version
                self.output_file = zip_path
                self.compression_time = time.time() - start_time  # <-- Store total time in seconds
                self.compression_time_str = f"{int(self.compression_time//60)}m {int(self.compression_time%60)}s"  # <-- Formatted string
                return zip_path
            
            return None

        except Exception as e:
            logging.error(f"Compression error: {e}")
            # If anything failed, try to clean up
            if os.path.exists(zip_path):
                try:
                    os.remove(zip_path)
                except:
                    pass
            return None

    def update_compression_ui(self, progress):
        """Update UI with compression progress"""
        self.root.after(0, lambda: [
            self.finalizing_progress_bar.set(progress),
            self.finalizing_time_label.configure(
                text=f"Compressing {progress*100:.1f}% (ZIP64)" if progress > 0.5 
                else f"Compressing {progress*100:.1f}%"
            )
        ])

    def update_compression_progress(self, progress):
        """Update UI during compression"""
        self.finalizing_progress_bar.set(progress)
        self.finalizing_time_label.configure(
            text=f"Compressing ({progress*100:.1f}%)"
        )

    def open_post_analysis_window(self):
        """Open the post-analysis window"""
        try:
            self.clear_container()
            
            # Frame for the final screen
            container = ctk.CTkFrame(self.main_container, fg_color="transparent")
            container.pack(expand=True, fill="both", padx=20, pady=20)

            # Title
            title = ctk.CTkLabel(
                container,
                text="Post-Analysis Actions",
                font=("Segoe UI", 22, "bold"),
                text_color=PRIMARY_COLOR
            )
            title.pack(pady=(10, 30))

            # Buttons for actions
            actions = [
                ("📂 Open Output Folder", lambda: os.startfile(os.path.dirname(self.output_file))),
                ("🛡️ Anti-Forensics Scan", lambda: self.check_for_anti_forensics(self.output_file))
            ]

            for label, command in actions:
                btn = ctk.CTkButton(
                    container,
                    text=label,
                    font=("Segoe UI", 16),
                    height=50,
                    width=300,
                    corner_radius=25,
                    fg_color=PRIMARY_COLOR,
                    hover_color=ACCENT_COLOR,
                    command=command
                )
                btn.pack(pady=10)

            # Bottom buttons: back home & re-capture
            bottom_frame = ctk.CTkFrame(container, fg_color="transparent")
            bottom_frame.pack(side="bottom", pady=30)

            back_btn = ctk.CTkButton(
                bottom_frame,
                text="⌂ Back to Home",
                font=("Segoe UI", 14),
                width=150,
                height=40,
                corner_radius=20,
                fg_color=PRIMARY_COLOR,
                hover_color=ACCENT_COLOR,
                command=self.show_welcome_screen
            )
            back_btn.pack(side="left", padx=20)

            recapture_btn = ctk.CTkButton(
                bottom_frame,
                text="🔁 Re-Capture",
                font=("Segoe UI", 14),
                width=150,
                height=40,
                corner_radius=20,
                fg_color=PRIMARY_COLOR,
                hover_color=ACCENT_COLOR,
                command=self.show_step3
            )
            recapture_btn.pack(side="left", padx=20)
            
        except Exception as e:
            logging.error(f"Error opening post-analysis window: {e}")
            self.show_welcome_screen()

        self.add_utility_buttons()

    def finalize_processing(self, start_time, total_time):
        """Finalize the processing and open post-analysis window"""
        try:
            # Hide finalizing progress UI
            if hasattr(self, 'finalizing_progress_bar'):
                self.finalizing_progress_bar.pack_forget()
            if hasattr(self, 'finalizing_progress_label'):
                self.finalizing_progress_label.pack_forget()
            if hasattr(self, 'finalizing_time_label'):
                self.finalizing_time_label.pack_forget()
            
            # Save last session data
            self.config['last_session'] = {
                'case_name': self.user_data.get('case_name', ''),
                'investigator_name': self.user_data.get('investigator_name', ''),
                'investigator_id': self.user_data.get('investigator_id', ''),
                'memory_id': self.user_data.get('memory_id', ''),
                'compression': self.enable_compression,
                'volatility': self.enable_volatility,
                'preview': self.enable_preview,
                'selected_plugins': self.selected_plugins if self.enable_volatility else [],
                'output_file': self.output_file,
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            save_config(self.config)
            
            # Generate report with total time
            hash_value = calculate_hash(self.output_file, self.user_data["hash_algorithm"])
            image_size = os.path.getsize(self.output_file)
            
            report_kwargs = {}
            if hasattr(self, 'volatility_time_str') and self.volatility_time_str:
               report_kwargs['volatility_time'] = self.volatility_time_str
            if hasattr(self, 'compression_time_str') and self.compression_time_str:
               report_kwargs['compression_time'] = self.compression_time_str          

  
            generate_report(
                self.output_file,
                self.user_data["investigator_name"],
                int(self.user_data["investigator_id"]),
                self.user_data["case_name"],
                int(self.user_data["case_id"]),
                int(self.user_data["memory_id"]),
                hash_value,
                image_size,
                total_time,
                self.user_data["report_format"],
                self.user_data["hash_algorithm"],
                **report_kwargs
            )
            success_msg = (
                f"RAM image captured successfully!\n"
                f"Image saved at: {self.output_file}\n"
                f"Report generated in {self.user_data['report_format']} format\n\n"
                )

            show_info(success_msg)
            # Directly show the post-analysis window
            self.open_post_analysis_window()
            
        except Exception as e:
            logging.error(f"Error during finalization: {e}")
            show_error(f"Finalization failed: {str(e)}")
            self.show_welcome_screen()


    
    def stop_capture(self):
        self.capture_active = False
        if self.capture_process:
            self.capture_process.terminate()
        
        if os.path.exists(self.output_file):
            file_size = os.path.getsize(self.output_file)
            if file_size > 0 and file_size < self.total_memory_size:
                if messagebox.askyesno("Partial Capture", "Save partial capture for resuming later?"):
                    self.partial_capture_path = self.output_file
                    self.resume_capture = True
                    self.capture_bytes_copied = file_size
    
    def finish_capture(self, stopped=False, success=False, output_file="", report_file="", volatility_report=None):
        self.capture_active = False
        self.capture_btn.configure(
            text="⚡",
            command=self.start_capture,
            state="normal"
        )
        
        if stopped:
            self.progress_label.configure(text="Capture stopped")
            show_info("Capture process was stopped")
        elif success:
            self.progress_label.configure(text="Capture complete!")
            self.progress_bar.set(1)
            self.time_left_label.configure(text="") 
            success_msg = (
                f"RAM image captured successfully!\n"
                f"Image saved at: {output_file}\n"
                f"Report saved at: {report_file}\n\n"
               )
            
            if volatility_report:
                success_msg += f"\nVolatility report: {volatility_report}"
            
            show_info(success_msg)
            self.resume_capture = False
            self.partial_capture_path = None

def create_gui():
    root = ctk.CTk()
    app = MemoryCaptureApp(root)
    root.mainloop()

if __name__ == "__main__":
    create_gui()